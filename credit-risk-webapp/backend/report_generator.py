"""
Module tạo báo cáo Word chuyên nghiệp từ kết quả phân tích tín dụng
"""

import os
from io import BytesIO
from datetime import datetime
import pandas as pd

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    WORD_OK = True
except Exception:
    WORD_OK = False


def generate_word_report(
    ratios_display: pd.DataFrame,
    pd_value: float,
    pd_label: str,
    ai_analysis: str,
    fig_bar=None,
    fig_radar=None,
    company_name: str = "KHÁCH HÀNG DOANH NGHIỆP",
    logo_path: str = None
) -> BytesIO:
    """
    Tạo báo cáo Word chuyên nghiệp từ kết quả phân tích tín dụng.

    Args:
        ratios_display: DataFrame chứa 14 chỉ số tài chính (index = tên chỉ số, column = giá trị)
        pd_value: Xác suất vỡ nợ (PD) dưới dạng số float (0-1) hoặc NaN
        pd_label: Nhãn dự đoán ("Default" hoặc "Non-Default")
        ai_analysis: Text phân tích từ AI
        fig_bar: Matplotlib figure của bar chart (optional)
        fig_radar: Matplotlib figure của radar chart (optional)
        company_name: Tên công ty
        logo_path: Đường dẫn đến logo (optional)

    Returns:
        BytesIO object chứa Word document
    """
    if not WORD_OK:
        raise Exception("Thiếu thư viện python-docx. Vui lòng cài đặt: pip install python-docx Pillow")

    # Tạo document mới
    doc = Document()

    # Cấu hình margin
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ===== 1. HEADER VỚI LOGO VÀ TIÊU ĐỀ =====
    # Thêm logo nếu có
    if logo_path and os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(2.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    # Tiêu đề chính
    title = doc.add_heading('BÁO CÁO ĐÁNH GIÁ RỦI RO TÍN DỤNG', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(194, 24, 91)  # #c2185b
    title_run.font.bold = True

    # Subtitle
    subtitle = doc.add_paragraph('Dự báo Xác suất Vỡ nợ KHDN (PD) & Phân tích AI Chuyên sâu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = RGBColor(255, 107, 157)  # #ff6b9d
    subtitle_run.font.bold = True

    # Thông tin thời gian
    date_info = doc.add_paragraph(f"Ngày xuất báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_info.runs[0]
    date_run.font.size = Pt(10)

    # Thông tin khách hàng
    company_info = doc.add_paragraph()
    company_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company_info.add_run(f"Tên khách hàng: {company_name}")
    company_run.font.size = Pt(11)
    company_run.font.bold = True

    doc.add_paragraph()  # Spacer

    # ===== 2. KẾT QUẢ DỰ BÁO PD =====
    heading1 = doc.add_heading('1. KẾT QUẢ DỰ BÁO XÁC SUẤT VỠ NỢ (PD)', level=1)
    heading1_run = heading1.runs[0]
    heading1_run.font.color.rgb = RGBColor(255, 107, 157)

    pd_para = doc.add_paragraph()
    if pd.notna(pd_value):
        pd_para.add_run(f"Xác suất Vỡ nợ (PD): ").bold = True
        pd_para.add_run(f"{pd_value:.2%}\n")
        pd_para.add_run("Phân loại: ").bold = True
        pd_para.add_run(f"{pd_label}\n")

        if "Default" in pd_label and "Non-Default" not in pd_label:
            risk_run = pd_para.add_run("⚠️ RỦI RO CAO - CẦN XEM XÉT KỸ LƯỠNG")
            risk_run.bold = True
            risk_run.font.color.rgb = RGBColor(220, 53, 69)  # Red
        else:
            safe_run = pd_para.add_run("✓ RỦI RO THẤP - KHẢ QUAN")
            safe_run.bold = True
            safe_run.font.color.rgb = RGBColor(40, 167, 69)  # Green
    else:
        pd_para.add_run("Xác suất Vỡ nợ (PD): ").bold = True
        pd_para.add_run("Không có dữ liệu")

    doc.add_paragraph()  # Spacer

    # ===== 3. BẢNG CHỈ SỐ TÀI CHÍNH =====
    heading2 = doc.add_heading('2. CHỈ SỐ TÀI CHÍNH CHI TIẾT', level=1)
    heading2_run = heading2.runs[0]
    heading2_run.font.color.rgb = RGBColor(255, 107, 157)

    # Tạo bảng
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Chỉ số Tài chính'
    hdr_cells[1].text = 'Giá trị'

    # Style header
    for cell in hdr_cells:
        cell_para = cell.paragraphs[0]
        cell_run = cell_para.runs[0]
        cell_run.font.bold = True
        cell_run.font.size = Pt(11)
        cell_run.font.color.rgb = RGBColor(255, 255, 255)
        # Set background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'FF6B9D')  # Pink
        cell._element.get_or_add_tcPr().append(shading_elm)
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for idx, row in ratios_display.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        value = row['Giá trị']
        row_cells[1].text = f"{value:.4f}" if pd.notna(value) else "N/A"
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()  # Spacer

    # ===== 4. BIỂU ĐỒ VISUALIZATION =====
    if fig_bar or fig_radar:
        doc.add_page_break()
        heading3 = doc.add_heading('3. TRỰC QUAN HÓA DỮ LIỆU', level=1)
        heading3_run = heading3.runs[0]
        heading3_run.font.color.rgb = RGBColor(255, 107, 157)

        # Bar chart
        if fig_bar:
            try:
                doc.add_heading('3.1. Biểu đồ Cột - Giá trị các Chỉ số', level=2)
                bar_buffer = BytesIO()
                fig_bar.savefig(bar_buffer, format='png', dpi=150, bbox_inches='tight')
                bar_buffer.seek(0)
                doc.add_picture(bar_buffer, width=Inches(6))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()  # Spacer
            except Exception as e:
                doc.add_paragraph(f"Không thể tạo biểu đồ cột: {str(e)}")

        # Radar chart
        if fig_radar:
            try:
                doc.add_heading('3.2. Biểu đồ Radar - Phân tích Đa chiều', level=2)
                radar_buffer = BytesIO()
                fig_radar.savefig(radar_buffer, format='png', dpi=150, bbox_inches='tight')
                radar_buffer.seek(0)
                doc.add_picture(radar_buffer, width=Inches(5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"Không thể tạo biểu đồ radar: {str(e)}")

    # ===== 5. PHÂN TÍCH AI =====
    doc.add_page_break()
    heading4 = doc.add_heading('4. PHÂN TÍCH AI & KHUYẾN NGHỊ TÍN DỤNG', level=1)
    heading4_run = heading4.runs[0]
    heading4_run.font.color.rgb = RGBColor(255, 107, 157)

    if ai_analysis and ai_analysis.strip():
        # Chia thành các đoạn và thêm vào document
        analysis_paragraphs = ai_analysis.split('\n')
        for para_text in analysis_paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text)
                # Highlight keywords
                if "CHO VAY" in para_text and "KHÔNG CHO VAY" not in para_text:
                    for run in para.runs:
                        if "CHO VAY" in run.text:
                            run.font.color.rgb = RGBColor(40, 167, 69)  # Green
                            run.bold = True
                elif "KHÔNG CHO VAY" in para_text:
                    for run in para.runs:
                        if "KHÔNG CHO VAY" in run.text:
                            run.font.color.rgb = RGBColor(220, 53, 69)  # Red
                            run.bold = True
    else:
        doc.add_paragraph("Chưa có phân tích từ AI.")

    # ===== 6. FOOTER =====
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Báo cáo này được tạo tự động bởi Hệ thống Đánh giá Rủi ro Tín dụng - Powered by AI & Machine Learning\n"
        f"© {datetime.now().year} Credit Risk Assessment System | Version 2.0 Premium"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.runs[0]
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(128, 128, 128)  # Grey

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
